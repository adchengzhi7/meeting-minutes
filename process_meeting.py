#!/usr/bin/env python3
"""
會議錄影/錄音 → Gemini 結構化會議記錄 → Google Doc
"""

import os
import sys
import subprocess
import tempfile
import time
from datetime import datetime
from pathlib import Path

import google.generativeai as genai
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
from docx import Document
from docx.shared import Pt
from dotenv import load_dotenv

# 所有路徑都基於腳本所在目錄，不依賴 working directory
PROJECT_DIR = Path(__file__).parent.resolve()
load_dotenv(PROJECT_DIR / ".env", override=True)

# ===== 設定 =====
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
GDRIVE_FOLDER_ID = os.getenv("GDRIVE_FOLDER_ID")

ARCHIVE_FOLDER = Path(os.getenv("ARCHIVE_FOLDER", "~/MeetingDrop/processed")).expanduser()
SCOPES = [
    "https://www.googleapis.com/auth/documents",
    "https://www.googleapis.com/auth/drive",
]

# 支援的檔案格式
SUPPORTED_VIDEO = {".mp4", ".mkv", ".mov", ".avi", ".webm"}
SUPPORTED_AUDIO = {".m4a", ".mp3", ".wav", ".ogg", ".flac", ".aac", ".wma"}
SUPPORTED_FORMATS = SUPPORTED_VIDEO | SUPPORTED_AUDIO

_gemini_configured = False


def _ensure_gemini():
    global _gemini_configured
    if not _gemini_configured:
        key = os.getenv("GEMINI_API_KEY")
        if not key:
            load_dotenv(PROJECT_DIR / ".env", override=True)
            key = os.getenv("GEMINI_API_KEY")
        genai.configure(api_key=key)
        _gemini_configured = True


def extract_audio(input_path: str, log=print) -> str:
    """從影片檔抽取音軌為 m4a"""
    suffix = Path(input_path).suffix.lower()

    if suffix in SUPPORTED_AUDIO:
        log("已是音訊檔，跳過抽取")
        return input_path

    log("正在抽取音軌...")
    output_path = tempfile.mktemp(suffix=".m4a")
    cmd = [
        "ffmpeg", "-i", input_path,
        "-vn",
        "-acodec", "copy",
        "-y",
        output_path
    ]

    result = subprocess.run(cmd, capture_output=True, text=True)

    if result.returncode != 0:
        log("音軌直接複製失敗，改用 AAC 編碼...")
        cmd = [
            "ffmpeg", "-i", input_path,
            "-vn",
            "-acodec", "aac",
            "-b:a", "128k",
            "-y",
            output_path
        ]
        subprocess.run(cmd, check=True, capture_output=True)

    log(f"音軌抽取完成")
    return output_path


def transcribe_and_summarize(audio_path: str, prompt_text: str, log=print, max_retries: int = 3) -> str:
    """用 Gemini 直接處理音檔，一步完成轉錄+摘要（含自動重試）"""
    _ensure_gemini()
    log("上傳音檔到 Gemini...")

    audio_file = genai.upload_file(audio_path)

    while audio_file.state.name == "PROCESSING":
        log("等待 Gemini 處理中...")
        time.sleep(5)
        audio_file = genai.get_file(audio_file.name)

    if audio_file.state.name == "FAILED":
        raise RuntimeError(f"Gemini 檔案處理失敗：{audio_file.state.name}")

    max_tokens = int(os.getenv("MAX_OUTPUT_TOKENS", "65536"))
    model = genai.GenerativeModel("gemini-2.5-flash")

    last_error = None
    for attempt in range(1, max_retries + 1):
        try:
            if attempt > 1:
                log(f"重試中（第 {attempt}/{max_retries} 次）...")
            else:
                log("呼叫 Gemini 產生會議記錄...")

            response = model.generate_content(
                [audio_file, prompt_text],
                generation_config=genai.GenerationConfig(
                    temperature=0.3,
                    max_output_tokens=max_tokens,
                ),
            )

            try:
                genai.delete_file(audio_file.name)
            except Exception:
                pass

            return response.text

        except Exception as e:
            last_error = e
            if attempt < max_retries:
                wait = 2 ** attempt
                log(f"Gemini API 錯誤：{e}，{wait} 秒後重試...")
                time.sleep(wait)
            else:
                try:
                    genai.delete_file(audio_file.name)
                except Exception:
                    pass
                raise RuntimeError(f"Gemini API 連續 {max_retries} 次失敗：{last_error}")


def get_google_creds():
    """OAuth：第一次跳瀏覽器授權，之後自動續用 token"""
    creds_path = PROJECT_DIR / "credentials.json"
    token_path = PROJECT_DIR / "token.json"
    creds = None

    if token_path.exists():
        creds = Credentials.from_authorized_user_file(str(token_path), SCOPES)

    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            flow = InstalledAppFlow.from_client_secrets_file(str(creds_path), SCOPES)
            creds = flow.run_local_server(port=0)

        token_path.write_text(creds.to_json())
        os.chmod(str(token_path), 0o600)

    return creds


def create_meeting_doc(title: str, markdown_content: str, folder_id: str = None) -> str:
    """建 .docx → 用使用者帳號上傳到 Drive → 回傳 Google Doc URL"""
    creds = get_google_creds()
    drive_service = build("drive", "v3", credentials=creds)

    docx_path = tempfile.mktemp(suffix=".docx")
    _markdown_to_docx(markdown_content, docx_path)

    try:
        file_metadata = {
            "name": title,
            "mimeType": "application/vnd.google-apps.document",
        }
        if folder_id:
            file_metadata["parents"] = [folder_id]

        media = MediaFileUpload(
            docx_path,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        doc_file = drive_service.files().create(
            body=file_metadata, media_body=media, fields="id",
        ).execute()
        doc_id = doc_file["id"]

        return f"https://docs.google.com/document/d/{doc_id}/edit"
    finally:
        Path(docx_path).unlink(missing_ok=True)


def _markdown_to_docx(markdown: str, output_path: str):
    """將 Markdown 轉成 .docx"""
    from docx.shared import Pt, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    doc = Document()

    # 調整預設字型（支援中文 fallback）
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)
    # 設定中文 fallback 字型
    rpr = style.element.get_or_add_rPr()
    ea_font = rpr.makeelement(qn("w:rFonts"), {})
    ea_font.set(qn("w:eastAsia"), "Microsoft JhengHei")
    rpr.insert(0, ea_font)

    # 調整標題樣式
    for level in range(1, 5):
        heading_style = doc.styles[f"Heading {level}"]
        heading_style.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
        if level <= 2:
            heading_style.paragraph_format.space_before = Pt(18)
            heading_style.paragraph_format.space_after = Pt(8)
        else:
            heading_style.paragraph_format.space_before = Pt(12)
            heading_style.paragraph_format.space_after = Pt(4)

    lines = markdown.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # 水平線
        if line.strip() in ("---", "***", "___"):
            _add_horizontal_rule(doc)

        # 標題
        elif line.startswith("#### "):
            doc.add_heading(line[5:], level=4)
        elif line.startswith("### "):
            doc.add_heading(line[3:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[2:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[1:].strip(), level=1)

        # 表格：收集連續的 | 行，轉成 docx table
        elif line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            i -= 1  # 回退一行，外層 i+=1 會補上
            _add_table(doc, table_lines)

        # 縮排列表（  - ）
        elif line.startswith("  - "):
            _add_rich_paragraph(doc, line.strip()[2:], style_name="List Bullet 2")

        # 列表
        elif line.startswith("- "):
            _add_rich_paragraph(doc, line[2:], style_name="List Bullet")

        # 空行跳過
        elif line.strip() == "":
            pass

        # 一般文字（支援粗體）
        else:
            _add_rich_paragraph(doc, line)

        i += 1

    doc.save(output_path)


def _add_horizontal_rule(doc):
    """加入水平線分隔"""
    from docx.oxml.ns import qn
    from docx.shared import Pt

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single",
        qn("w:sz"): "6",
        qn("w:space"): "1",
        qn("w:color"): "D0D0D0",
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_table(doc, table_lines: list):
    """將 Markdown 表格轉成 docx 表格"""
    from docx.shared import Pt, RGBColor, Cm
    from docx.oxml.ns import qn

    # 解析表格行
    rows = []
    for line in table_lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        # 跳過分隔行 |---|---|
        if cells and all(set(c) <= {"-", ":", " "} for c in cells):
            continue
        rows.append(cells)

    if not rows:
        return

    n_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Light Grid Accent 1"

    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            if c_idx < n_cols:
                cell = table.cell(r_idx, c_idx)
                cell.text = ""
                p = cell.paragraphs[0]
                _add_runs(p, cell_text)
                p.paragraph_format.space_after = Pt(2)
                # 表格內文字稍小
                for run in p.runs:
                    run.font.size = Pt(10)

                # 表頭粗體 + 底色
                if r_idx == 0:
                    for run in p.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    # 表頭底色（深藍）
                    tc = cell._element
                    tcPr = tc.get_or_add_tcPr()
                    shading = tcPr.makeelement(qn("w:shd"), {
                        qn("w:val"): "clear",
                        qn("w:color"): "auto",
                        qn("w:fill"): "2563EB",
                    })
                    tcPr.append(shading)

    doc.add_paragraph()  # 表格後空行


def _add_rich_paragraph(doc, text: str, style_name: str = None):
    """加入段落，支援 **粗體** 格式"""
    if style_name:
        p = doc.add_paragraph(style=style_name)
    else:
        p = doc.add_paragraph()
    _add_runs(p, text)


def _add_runs(paragraph, text: str):
    """將文字加入段落，解析 **粗體** 標記"""
    parts = text.split("**")
    for i, part in enumerate(parts):
        if part:
            run = paragraph.add_run(part)
            if i % 2 == 1:
                run.bold = True


def _save_history(original_name: str, doc_title: str, doc_url: str):
    """儲存處理記錄到 history.json"""
    import json as _json
    history_path = PROJECT_DIR / "history.json"

    records = []
    if history_path.exists():
        try:
            records = _json.loads(history_path.read_text())
        except Exception:
            records = []

    records.insert(0, {
        "original": original_name,
        "title": doc_title,
        "url": doc_url,
        "time": datetime.now().strftime("%Y-%m-%d %H:%M"),
    })

    # 最多保留 200 筆
    history_path.write_text(_json.dumps(records[:200], ensure_ascii=False, indent=2))


def send_notification(title: str, message: str, url: str = None):
    """macOS 系統通知"""
    script = f'display notification "{message}" with title "{title}" sound name "Glass"'
    subprocess.run(["osascript", "-e", script])

    if url:
        subprocess.run(["open", url])


def process_file(file_path: str, auto_open: bool = True, log=print):
    """處理單一檔案的完整流程"""
    file_path = Path(file_path).resolve()
    filename = file_path.stem
    suffix = file_path.suffix.lower()

    if suffix not in SUPPORTED_FORMATS:
        log(f"不支援的格式：{suffix}")
        return None

    log(f"開始處理：{file_path.name}")

    audio_path = extract_audio(str(file_path), log=log)
    temp_audio = audio_path != str(file_path)

    try:
        prompt_path = Path(__file__).parent / "prompt.md"
        prompt_text = prompt_path.read_text(encoding="utf-8")

        meeting_notes = transcribe_and_summarize(audio_path, prompt_text, log=log)

        # 從回覆第一行抓 AI 摘要標題
        lines = meeting_notes.strip().split("\n")
        ai_title = lines[0].strip().lstrip("#").strip() if lines else filename
        # 移除第一行（標題），剩餘為正文
        doc_body = "\n".join(lines[1:]).strip()

        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M")
        doc_title = f"{timestamp} — {ai_title}"
        log("建立 Google Doc...")
        doc_url = create_meeting_doc(doc_title, doc_body, GDRIVE_FOLDER_ID or None)

        ARCHIVE_FOLDER.mkdir(parents=True, exist_ok=True)
        archive_path = ARCHIVE_FOLDER / file_path.name
        file_path.rename(archive_path)
        log(f"原始檔案已歸檔")

        # 儲存 metadata
        _save_history(filename, doc_title, doc_url)

        send_notification(
            "會議記錄完成",
            f"{filename} 的會議記錄已產生",
            doc_url if auto_open else None,
        )

        log(f"完成！{doc_url}")
        return doc_url

    finally:
        if temp_audio and os.path.exists(audio_path):
            os.remove(audio_path)


def main():
    if len(sys.argv) < 2:
        print("用法：python process_meeting.py <錄影/錄音檔路徑>")
        print(f"支援格式：{', '.join(sorted(SUPPORTED_FORMATS))}")
        sys.exit(1)

    file_path = sys.argv[1]
    if not os.path.exists(file_path):
        print(f"檔案不存在：{file_path}")
        sys.exit(1)

    process_file(file_path)


if __name__ == "__main__":
    main()
